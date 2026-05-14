import streamlit as st
import requests
from docx import Document
from io import BytesIO

# Page configuration
st.set_page_config(
    page_title="AI Data Assistant",
    page_icon="📊",
    layout="wide" 
)

# Style the app with custom CSS
st.markdown(
    """
<style>
    [data-testid="stAppViewContainer"]
        { background-color: #F0F2F6; }
    .stAppDeployButton { 
        display: none; }
    header { 
        visibility: hidden; }
    [data-testid="stSidebar"] 
        { background-color: #2E7BCF; color: #ffffff; }
    [data-testid="stSidebar"] 
    [data-testid="stImage"] { display: block; margin-left: auto; margin-right: auto; }
    [data-testid="stSidebar"] .stMarkdown { text-align: center; }
    [data-testid="stWidgetLabel"] p { color: white; font-weight: bold; }
    [data-testid="stFileUploaderDropzone"] button {
        text-indent: -9999px; line-height: 0; background-color: #2E7BCF;
        color: #ffffff; border: none; padding: 13px 20px; border-radius: 5px;
        }
    [data-testid="stFileUploaderDropzone"] button::after { content: "Upload File"; text-indent: 0; line-height: initial; display: block; }

   .stDownloadButton > button {
        width: 100%;
        border-radius: 8px;
        color: #ffffff;
        background-color: #4CAF50; /* Modern Green */
        border: none;
        padding: 0.6rem;
        transition: all 0.3s ease;
        font-weight: 600;
    }
    .stDownloadButton > button:hover {
        background-color: #45a049;
        border: none;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        transform: translateY(-1px);
        color: #ffffff;
    }
    .stChatMessage {
        border: 2px solid #cbd5e1 !important; 
        border-radius: 8px !important;
        background-color: #ffffff !important;
        color: #1e293b !important; }
    .stChatMessage.user { background-color: #DCF8C6; align-self: flex-end; }
    .stChatMessage.assistant { background-color: #ffffff; align-self: flex-start; }
        
    [data-testid="stChatInput"] > div {
        border: 1px solid #1E6091 !important; /* Forces your blue color */
        border-radius: 20px !important;
        padding-left: 10px !important;
        height: 60px;
        background-color: white !important;
    }

    [data-testid="stChatInput"] div[role="textbox"] {
        border: none !important;
        box-shadow: none !important;
    }

    [data-testid="stChatInput"] :focus-within {
        border-color: #1E6091 !important;
        box-shadow: none !important;
        outline: none !important;
        border: none !important;
    }

    [data-testid="stChatInput"] textarea {
        border: none !important;
        box-shadow: none !important;
        outline: none !important;
    }

    [data-testid="stChatInput"] button {
        right: 10px !important;
        background-color: transparent !important;
        align-self: center !important;
    }
    /* 4. Improve overall font and spacing */
    h1, h2, h3 {
        font-family: 'Inter', sans-serif;
        color: #1E1E1E;
    }
    </style>""", unsafe_allow_html=True
)

# --- Helper Function for Word Export ---
def generate_docx(messages):
    doc = Document()
    doc.add_heading('AI Data Analysis Report', 0)
    
    for msg in messages:
        role = "User" if msg["role"] == "user" else "AI Assistant"
        p = doc.add_paragraph()
        p.add_run(f"{role}: ").bold = True
        p.add_run(msg["content"])
    
    # Save to a buffer so we don't need to save a real file on the server
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- SIDEBAR: LOGO & UPLOAD ---
with st.sidebar:
    st.image("logo.png", width=120)
    st.markdown("Your AI Data Assistant")
    st.markdown("---")
    uploaded_file = st.file_uploader("Upload CSV or Excel file", type=["csv", "xlsx", "xls"])

# Check if there are any messages to download
if "messages" in st.session_state and st.session_state.messages:
    st.sidebar.markdown("---")
    st.sidebar.subheader("Export Results")
    
    # Generate the Word file
    docx_file = generate_docx(st.session_state.messages)
    
    st.sidebar.download_button(
        label="📄 Download Report (.docx)",
        data=docx_file,
        file_name="analysis_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    st.markdown("---")

# --- MAIN PAGE: PROJECT NAME ---
st.title(" 📊 AI Data Assistant")
st.markdown("##### *Transforming raw data into deterministic business insights.*")

# --- CHAT INITIALIZATION ---
# 1. Initialize chat history in session state if it doesn't exist
if "messages" not in st.session_state:
    st.session_state.messages = []

# 2. Display all previous messages from history
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# 3. Chat Input Logic
if prompt := st.chat_input("Ask about your data..."):
    # Add user message to UI and history
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # Prepare data for Backend
    with st.spinner("Analyzing..."):
        # Format history string for the AI's context
        history_context = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages[:-1]])
        
        files = {"file": (uploaded_file.name, uploaded_file.getvalue())}
        data = {"prompt": prompt, "history": history_context}

        try:
            response = requests.post("http://127.0.0.1:8000/analyze", files=files, data=data)
            
            if response.status_code == 200:
                insight = response.json()["insight"]
                
                # Add assistant response to UI and history
                with st.chat_message("assistant"):
                    st.markdown(insight)
                st.session_state.messages.append({"role": "assistant", "content": insight})
            else:
                st.error(f"Backend Error: {response.text}")
        except Exception as e:
            st.error(f"Connection Error: {e}")
else:
    # Welcome screen before file upload
    st.markdown("""
    ### Hi 👋 , Welcome to your AI Data Assistant.
    1. **Upload** your File using the sidebar on the left.
    2. **Chat** with the AI to get deterministic insights.
    """)